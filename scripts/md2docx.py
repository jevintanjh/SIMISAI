#!/usr/bin/env python3
"""
Permanent markdown to .docx converter for SIMISAI
Usage: python3 scripts/md2docx.py <input.md> [output.docx]
If output not specified, creates .docx with same name as input
"""

import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor
import re

def add_heading(doc, text, level):
    """Add a heading with appropriate styling"""
    heading = doc.add_heading(text, level=level)
    return heading

def process_markdown_line(doc, line):
    """Process a single markdown line and add to document"""
    line = line.rstrip()

    # Skip empty lines
    if not line:
        doc.add_paragraph()
        return

    # Horizontal rule
    if line.strip() == '---':
        doc.add_paragraph('_' * 50)
        return

    # Headers
    if line.startswith('#'):
        match = re.match(r'^(#{1,6})\s+(.+?)(?:\s+#{1,6})?$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip('*')
            add_heading(doc, text, level)
            return

    # Bullet lists
    if line.startswith('- '):
        text = line[2:]
        text = format_inline_markdown(text)
        para = doc.add_paragraph(style='List Bullet')
        add_formatted_text(para, text)
        return

    # Numbered lists
    match = re.match(r'^(\d+)\.\s+(.+)$', line)
    if match:
        text = match.group(2)
        text = format_inline_markdown(text)
        para = doc.add_paragraph(style='List Number')
        add_formatted_text(para, text)
        return

    # Regular paragraph
    text = format_inline_markdown(line)
    para = doc.add_paragraph()
    add_formatted_text(para, text)

def format_inline_markdown(text):
    """Keep markdown formatting markers for later processing"""
    return text

def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, links)"""
    # Process bold text (**text**)
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Check for links [text](url)
            link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
            if '[' in part and '](' in part:
                # Simple link handling - just show the text
                part = re.sub(link_pattern, r'\1', part)
            paragraph.add_run(part)

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to .docx"""
    # Check if input file exists
    if not os.path.exists(md_file):
        print(f"Error: Input file '{md_file}' not found")
        sys.exit(1)

    # Create document
    doc = Document()

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Process each line
    for line in lines:
        process_markdown_line(doc, line)

    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/md2docx.py <input.md> [output.docx]")
        print("Example: python3 scripts/md2docx.py docs/investor/case-study.md")
        sys.exit(1)

    input_file = sys.argv[1]

    # If output file not specified, use same name with .docx extension
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        output_file = os.path.splitext(input_file)[0] + '.docx'

    convert_markdown_to_docx(input_file, output_file)
